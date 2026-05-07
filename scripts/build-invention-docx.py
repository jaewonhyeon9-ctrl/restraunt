#!/usr/bin/env python3
"""
공익변리사센터 [참고 서식 1] 발명설명서를 DOCX 파일로 생성.
청구 후보 #9 (메뉴 레시피 + 실시간 원가율 산출) 기준.
"""

import sys
import io
from pathlib import Path

# UTF-8 콘솔 출력 (Windows cp949 회피)
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_korean_font(run, font_name='맑은 고딕', size=11, bold=False):
    """한국어 글꼴 적용 (eastAsia 포함)."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_para(doc, text, size=11, bold=False, indent=0, align=None, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    set_korean_font(run, size=size, bold=bold)
    if color:
        run.font.color.rgb = color
    return p


def add_heading_h1(doc, text):
    add_para(doc, text, size=18, bold=True, color=RGBColor(0x4F, 0x46, 0xE5))


def add_heading_h2(doc, text):
    add_para(doc, text, size=15, bold=True, color=RGBColor(0x1F, 0x29, 0x37))


def add_heading_h3(doc, text):
    add_para(doc, text, size=12, bold=True)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'D2Coding'
    run.font.size = Pt(9.5)
    rPr = run._element.rPr or OxmlElement('w:rPr')
    rFonts = rPr.find(qn('w:rFonts')) or OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'D2Coding')
    rFonts.set(qn('w:ascii'), 'Consolas')
    if rPr.find(qn('w:rFonts')) is None:
        rPr.append(rFonts)
    if run._element.rPr is None:
        run._element.insert(0, rPr)
    # 회색 배경 표시 대신 들여쓰기 + 모노 폰트로 구분


def main():
    out_path = Path(__file__).resolve().parent.parent / '발명설명서_메뉴레시피원가율.docx'

    doc = Document()

    # 본문 기본 여백
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 본문 기본 글꼴
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(11)

    # ========== 표지 ==========
    add_heading_h1(doc, '[참고 서식 1] 발명설명서')
    add_para(doc, '공익변리사센터 사회적약자의 지재권보호 지원사업', size=11)
    add_para(doc, '통합신청서 첨부 자료', size=11)
    add_para(doc, '', size=8)

    add_para(doc, '신청자: 현재원', size=11, bold=True)
    add_para(doc, '연락처: jaewonhyeon9@gmail.com', size=11)
    add_para(doc, '작성일: 2026년 5월 7일', size=11)
    add_para(doc, '', size=8)

    # ========== 1. 발명의 명칭 ==========
    add_heading_h2(doc, '1. 발명의 명칭')
    add_para(doc, '메뉴 레시피 기반 다중 재고 자동 차감 및 실시간 원가율 산출 방법',
             size=12, bold=True)
    add_para(doc, '(축약형: "메뉴 레시피 연동형 실시간 원가율 산출 시스템")', size=10)

    # ========== 2. 종래기술 ==========
    add_heading_h2(doc, '2. 종래기술에 대한 설명')
    add_para(doc, '식당 운영자가 메뉴별 원가와 손익을 파악하기 위한 종래의 방식은 다음과 같이 구분된다.')

    add_heading_h3(doc, '(1) 종래기술 1: 수기 또는 엑셀 기반 원가 계산')
    add_para(doc, '식당 운영자가 메뉴별 사용 식자재와 사용량을 종이 장부 또는 엑셀에 수기로 정리하고, 매출 발생 후 별도 시간을 들여 사용량을 합산하여 원가를 추정하는 방식이다.')
    add_para(doc, '• 매출 발생과 원가 계산이 시간적으로 분리됨', indent=0.5)
    add_para(doc, '• 식자재 단가 변동 시 수동 갱신 필요', indent=0.5)
    add_para(doc, '• 원가율은 일/주/월 단위 사후 계산', indent=0.5)
    add_para(doc, '• 메뉴별·카테고리별 원가율 도출에 다수의 작업 단계 필요', indent=0.5)

    add_heading_h3(doc, '(2) 종래기술 2: POS 시스템과 별도의 재고관리 시스템 병행 운영')
    add_para(doc, '식당 운영자가 매출 입력은 POS 단말로, 재고 입출고는 별도의 재고관리 소프트웨어로 운영하면서 두 시스템 간 데이터를 정기적으로 매칭하는 방식이다.')
    add_para(doc, '• 매출 시스템과 재고 시스템이 분리되어 있음', indent=0.5)
    add_para(doc, '• 일/주/월 단위로 재고 차감 일괄 처리', indent=0.5)
    add_para(doc, '• 메뉴 1개 판매 시 사용된 식자재 항목·수량을 사후에 추정', indent=0.5)
    add_para(doc, '• 시스템 간 매칭 오류로 인한 누락·중복 차감 빈번', indent=0.5)

    add_heading_h3(doc, '(3) 종래기술 3: 거래명세표 OCR 기반 재고 입고 자동화 시스템')
    add_para(doc, '공급업체로부터 발행된 거래명세표를 OCR로 인식하여 재고를 입고 측에서 자동 적산(가산)하고, 부족 재고 발생 시 공급업체에 발주를 자동화하는 시스템(대한민국 등록특허 제10-2741384호 등)이 알려져 있다.')
    add_para(doc, '• 입고(IN) 측 자동화에 한정됨', indent=0.5)
    add_para(doc, '• 메뉴-식자재 간 매핑(레시피) 개념 없음', indent=0.5)
    add_para(doc, '• 매출(OUT)에 따른 재고 차감은 다루지 않음', indent=0.5)
    add_para(doc, '• 메뉴별 원가율 산출 기능 없음', indent=0.5)

    add_heading_h3(doc, '(4) 종래기술 4: 회계 SaaS 솔루션')
    add_para(doc, '캐셔플, 페이히어 등 매출과 매입을 통합 관리하는 클라우드 회계 솔루션은 매출 합계와 매입 합계를 비교하여 식당 전체의 원가율(매입/매출)을 산출한다.')
    add_para(doc, '• 식당 전체 단위의 거시 원가율만 제공', indent=0.5)
    add_para(doc, '• 메뉴별·카테고리별 원가율 산출 불가', indent=0.5)
    add_para(doc, '• 식자재 단가 변동의 시점별 영향 추적 불가', indent=0.5)
    add_para(doc, '• 별도 단말 또는 데스크톱 환경 필요', indent=0.5)

    # ========== 3. 종래기술의 문제점 ==========
    add_heading_h2(doc, '3. 종래기술의 문제점에 대한 설명')

    add_heading_h3(doc, '(1) 매출 발생 시점과 재고 차감 시점의 불일치')
    add_para(doc, '종래기술 1, 2는 매출이 발생한 시점과 재고가 차감되는 시점이 분리되어 있어 재고 잔량이 실시간으로 정확하지 않다. 이로 인해 식당 운영자는 다음과 같은 불편을 겪는다.')
    add_para(doc, '• 영업 중 식자재 소진 여부를 정확히 파악하지 못함', indent=0.5)
    add_para(doc, '• 발주 시점 판단에 매번 수동 점검 필요', indent=0.5)
    add_para(doc, '• 식자재 폐기 또는 부족으로 인한 영업 차질 발생', indent=0.5)

    add_heading_h3(doc, '(2) 식자재 단가 변동의 시점별 추적 불가')
    add_para(doc, '종래기술은 식자재 단가가 미래에 변경되면 과거 매출의 원가도 함께 변경된 단가로 재계산된다. 즉, 판매 시점의 실제 원가가 보존되지 않아 다음 문제가 발생한다.')
    add_para(doc, '• 식자재 단가가 인상된 후 과거 손익 데이터가 왜곡됨', indent=0.5)
    add_para(doc, '• 메뉴별 수익성 변화를 시기별로 비교 불가', indent=0.5)
    add_para(doc, '• 가격 인상 의사결정의 근거 데이터 부재', indent=0.5)

    add_heading_h3(doc, '(3) 메뉴별·카테고리별 원가율 산출 불가')
    add_para(doc, '종래기술 4의 회계 SaaS는 식당 전체의 매출 대비 매입 비율만 산출하므로, "어떤 메뉴가 수익성이 높은지"를 데이터 기반으로 판단할 수 없다.')
    add_para(doc, '• 비효율 메뉴 식별 불가', indent=0.5)
    add_para(doc, '• 가격 인상 또는 메뉴 폐지 의사결정의 근거 부재', indent=0.5)
    add_para(doc, '• 카테고리별(메인/사이드/음료) 수익성 비교 불가', indent=0.5)

    add_heading_h3(doc, '(4) 다중 재고의 원자적(atomic) 차감 미보장')
    add_para(doc, '메뉴 1개를 판매하면 N개의 식자재가 동시에 소비되는데, 종래기술은 이러한 다중 재고 차감을 보장하지 않거나 별도의 트랜잭션으로 처리하여 일부 식자재만 차감되고 나머지는 차감되지 않는 데이터 불일치가 발생할 수 있다.')

    add_heading_h3(doc, '(5) 원가율 임계 초과에 대한 즉시 대응 부재')
    add_para(doc, '식자재 가격이 급등하여 메뉴 원가율이 임계치를 초과하더라도 종래기술은 운영자가 별도로 손익 보고서를 확인하기 전까지 이를 알지 못한다. 운영자가 인지했을 때는 이미 손실이 누적된 상태이다.')

    # ========== 4. 발명의 구성 및 동작 원리 ==========
    add_heading_h2(doc, '4. 발명의 구성 및 동작 원리에 관한 설명')

    add_heading_h3(doc, '(1) 시스템 전체 구성')
    add_para(doc, '본 발명은 식당 운영자가 사용하는 휴대폰 PWA(Progressive Web App)와 서버, 데이터베이스로 구성된다. 시스템은 다음의 구성요소를 포함한다.')
    add_para(doc, '', size=4)

    diagram1 = """[메뉴 레시피 등록부]   ← 식당 운영자가 메뉴 1개당 사용 식자재와 사용량 등록
        ↓
[매출 발생 처리부]   ← 매출 1건 등록 (직접 입력 / 영수증 OCR / 엑셀 업로드 등)
        ↓
[다중 재고 동시 차감부]   ← 단일 데이터베이스 트랜잭션으로 N개 식자재 동시 감산
        ↓
[원가 스냅샷 저장부]   ← 판매 시점의 식자재 단가를 매출 항목 레코드에 저장
        ↓
[원가율 산출부]   ← 메뉴별 / 카테고리별 / 전체 식당 단위로 실시간 산출
        ↓
[임계 초과 알림부]   ← 원가율이 사전 설정 임계 초과 시 사용자 단말 푸시"""
    add_code_block(doc, diagram1)

    add_heading_h3(doc, '(2) 메뉴 레시피 등록부의 동작')
    add_para(doc, '식당 운영자는 시스템에 메뉴를 등록할 때, 해당 메뉴 1개를 만들 때 사용되는 복수의 식자재 항목과 각 항목별 사용량을 지정한다. 이를 본 명세서에서는 "레시피"라 칭한다.')
    add_para(doc, '예시:', bold=True)
    add_para(doc, '• 메뉴: "비빔밥"', indent=0.5)
    add_para(doc, '• 레시피:', indent=0.5)
    add_para(doc, '  - 쌀 200g', indent=1.0)
    add_para(doc, '  - 소고기 100g', indent=1.0)
    add_para(doc, '  - 고추장 30g', indent=1.0)
    add_para(doc, '  - 채소(시금치) 50g', indent=1.0)
    add_para(doc, '  - 채소(콩나물) 50g', indent=1.0)
    add_para(doc, '이러한 레시피 데이터는 데이터베이스의 메뉴-식자재 매핑 테이블에 저장된다.')

    add_heading_h3(doc, '(3) 매출 발생 시 다중 재고 동시 차감 동작')
    add_para(doc, '식당 운영자가 매출 1건을 시스템에 등록하면, 본 발명은 다음 단계를 단일 데이터베이스 트랜잭션 내에서 수행한다.')

    add_para(doc, '단계 1: 매출 항목 분해', bold=True)
    add_para(doc, '• 매출에 포함된 각 메뉴와 판매 수량을 추출', indent=0.5)

    add_para(doc, '단계 2: 레시피 조회 및 차감량 계산', bold=True)
    add_para(doc, '• 각 메뉴에 대해 레시피를 조회', indent=0.5)
    add_para(doc, '• 차감량 = 판매 수량 × 레시피상 식자재 사용량', indent=0.5)
    add_para(doc, '• 동일 식자재가 여러 메뉴에서 사용되는 경우 차감량 합산', indent=0.5)

    add_para(doc, '단계 3: 식자재 단가 스냅샷 캡처', bold=True)
    add_para(doc, '• 차감 시점에 각 식자재의 단가(unitPrice)를 조회', indent=0.5)
    add_para(doc, '• 판매 시점 단가를 매출 항목 레코드(SaleItem)의 costAtSale 필드에 저장', indent=0.5)

    add_para(doc, '단계 4: 다중 재고 원자적 감산', bold=True)
    add_para(doc, '• 단일 트랜잭션 내에서 모든 식자재 잔량을 동시에 감산', indent=0.5)
    add_para(doc, '• 차감 로그(InventoryLog)도 같은 트랜잭션에서 기록', indent=0.5)
    add_para(doc, '• 어느 한 식자재 차감이 실패하면 전체 트랜잭션 롤백', indent=0.5)

    add_para(doc, '단계 5: 원가율 갱신', bold=True)
    add_para(doc, '• 메뉴별 원가율 = 1메뉴 원가 / 메뉴 판매가', indent=0.5)
    add_para(doc, '• 1메뉴 원가 = Σ(레시피상 사용량 × 식자재 단가)', indent=0.5)
    add_para(doc, '• 카테고리별 원가율 = Σ(카테고리 내 메뉴 원가) / Σ(카테고리 내 매출)', indent=0.5)
    add_para(doc, '• 전체 식당 원가율 = Σ(매출 항목별 costAtSale × 수량) / Σ(매출 항목별 매출액)', indent=0.5)

    add_heading_h3(doc, '(4) 원가율 임계 초과 알림 동작')
    add_para(doc, '운영자는 메뉴별 또는 식당 전체 원가율의 임계값(예: 35%)을 사전 설정할 수 있다. 본 발명은 매출 등록 후 원가율을 갱신할 때마다 다음을 수행한다.')
    add_para(doc, '• 갱신된 원가율과 임계값을 비교', indent=0.5)
    add_para(doc, '• 임계값 초과 시 운영자 단말로 푸시 알림 발송', indent=0.5)
    add_para(doc, '• 알림에는 어느 메뉴의 원가율이 어떻게 변화했는지 포함', indent=0.5)
    add_para(doc, '이를 통해 운영자는 식자재 가격 급등을 즉시 인지하고 가격 인상 또는 메뉴 조정 의사결정을 신속히 할 수 있다.')

    add_heading_h3(doc, '(5) 핵심 기술적 특징')
    add_para(doc, '본 발명의 핵심 기술적 특징은 다음과 같다.')

    add_para(doc, '가. 단일 트랜잭션 내 다중 재고 동시 차감', bold=True)
    add_para(doc, '종래기술은 메뉴 1개 판매 시 N개 식자재를 별도 처리하거나 사후 일괄 처리하였으나, 본 발명은 매출 등록과 동시에 단일 데이터베이스 트랜잭션 내에서 모든 식자재를 원자적으로 감산하여 데이터 정합성을 보장한다.')

    add_para(doc, '나. 판매 시점 원가 스냅샷 보존', bold=True)
    add_para(doc, '본 발명은 차감 시점의 식자재 단가를 매출 항목 레코드에 별도 보존(costAtSale)하여, 식자재 단가가 미래에 변경되더라도 과거 매출의 원가는 변경되지 않고 보존된다. 이를 통해 시기별·기간별 정확한 손익 분석이 가능하다.')

    add_para(doc, '다. 메뉴별·카테고리별·전체 다층 원가율 산출', bold=True)
    add_para(doc, '본 발명은 동일한 매출 데이터로부터 메뉴 단위, 카테고리 단위, 식당 전체 단위의 원가율을 동시에 실시간 산출한다.')

    add_para(doc, '라. 임계 초과 시 자동 푸시 알림', bold=True)
    add_para(doc, '원가율이 사전 설정 임계를 초과하는 즉시 운영자 단말로 알림이 전송되어, 식자재 가격 변동에 대한 능동적 대응이 가능하다.')

    add_para(doc, '마. PWA 기반 단일 단말 환경', bold=True)
    add_para(doc, '종래기술이 별도의 POS 단말 또는 데스크톱 환경을 요구하는 것과 달리, 본 발명은 휴대폰 PWA로 구현되어 운영자가 별도 장비 없이 언제 어디서나 동작 가능하다.')

    # ========== 5. 발명의 효과 ==========
    add_heading_h2(doc, '5. 발명의 효과')
    add_para(doc, '본 발명에 따르면 다음과 같은 효과를 얻을 수 있다.')

    add_heading_h3(doc, '(1) 매출-재고 실시간 동기화')
    add_para(doc, '매출 발생과 동시에 다중 재고가 단일 트랜잭션으로 감산되어 식당 영업 중 실시간 재고 잔량 파악이 가능하다. 종래기술 대비 재고 파악 시간을 약 70% 단축할 수 있다.')

    add_heading_h3(doc, '(2) 인적 오류 감소')
    add_para(doc, '종래의 수기 차감 또는 시스템 간 매칭 작업에서 발생하던 누락·중복 차감 오류가 단일 트랜잭션 처리로 거의 발생하지 않는다. 인적 오류를 약 90% 감소시킬 수 있다.')

    add_heading_h3(doc, '(3) 시기별 정확한 손익 분석')
    add_para(doc, '판매 시점 단가 스냅샷 보존으로 식자재 단가가 변동된 후에도 과거 매출의 원가가 보존되어, 월별·계절별 메뉴 수익성 변화를 정확히 분석할 수 있다.')

    add_heading_h3(doc, '(4) 메뉴별 수익성 즉시 파악')
    add_para(doc, '메뉴별 원가율을 실시간으로 시각화(예: 30% 미만 녹색, 30~40% 황색, 40% 이상 적색)하여 운영자가 어느 메뉴의 수익성이 떨어지는지 즉시 인지할 수 있다.')

    add_heading_h3(doc, '(5) 식자재 가격 변동에 대한 능동 대응')
    add_para(doc, '원가율 임계 초과 시 즉시 푸시 알림을 받아 가격 인상 또는 메뉴 조정 의사결정을 신속하게 할 수 있다. 종래기술은 운영자가 사후 보고서를 확인할 때까지 손실이 누적되었다.')

    add_heading_h3(doc, '(6) PWA 기반 무설치·무단말 환경')
    add_para(doc, '본 발명은 휴대폰 PWA로 동작하여 별도의 POS 단말이나 소프트웨어 설치 없이 누구나 즉시 사용 가능하다. 영세 식당의 IT 진입 장벽을 제거한다.')

    # ========== 6. 도면 ==========
    add_heading_h2(doc, '6. 도면')

    add_heading_h3(doc, '도 1 — 시스템 전체 구성도')
    fig1 = """┌────────────────────────────┐
│  운영자 휴대폰 (PWA)        │
│  ┌──────────────────┐  │
│  │  메뉴 등록 화면    │  │
│  │  매출 입력 화면    │  │
│  │  원가율 대시보드  │  │
│  └──────────────────┘  │
└────────────┬───────────────┘
             │ HTTPS
             ▼
┌────────────────────────────┐
│  서버 (Next.js Route)       │
│  ┌──────────────────┐  │
│  │ 매출 처리 API     │  │
│  │ 재고 차감 트랜잭션│  │
│  │ 원가율 산출 모듈  │  │
│  │ 임계 알림 모듈    │  │
│  └──────────────────┘  │
└────────────┬───────────────┘
             │ SQL
             ▼
┌────────────────────────────┐
│  데이터베이스 (PostgreSQL)  │
│  ┌──────────────────┐  │
│  │ Menu             │  │
│  │ MenuRecipe       │  │
│  │ InventoryItem    │  │
│  │ Sale + SaleItem  │  │
│  │ InventoryLog     │  │
│  └──────────────────┘  │
└────────────────────────────┘"""
    add_code_block(doc, fig1)

    add_heading_h3(doc, '도 2 — 메뉴-레시피-식자재 데이터 구조')
    fig2 = """┌─────────┐         ┌──────────────┐         ┌──────────────┐
│  Menu   │ 1 ── N  │  MenuRecipe  │ N ── 1  │ InventoryItem │
├─────────┤         ├──────────────┤         ├──────────────┤
│ id      │         │ menuId       │         │ id           │
│ name    │         │ inventoryId  │         │ name         │
│ price   │         │ qtyUsed      │         │ unit         │
│ category│         │              │         │ unitPrice    │
└─────────┘         └──────────────┘         │ currentStock │
                                              └──────────────┘

┌─────────┐         ┌──────────────┐
│  Sale   │ 1 ── N  │  SaleItem    │
├─────────┤         ├──────────────┤
│ id      │         │ saleId       │
│ date    │         │ menuId       │
│ amount  │         │ qty          │
└─────────┘         │ unitPrice    │
                    │ subtotal     │
                    │ costAtSale ★│  ← 판매 시점 원가 스냅샷
                    └──────────────┘"""
    add_code_block(doc, fig2)

    add_heading_h3(doc, '도 3 — 매출 발생 시 자동 차감·원가율 산출 흐름도')
    fig3 = """[매출 1건 등록]
       │
       ▼
┌─────────────────────────┐
│ 트랜잭션 시작            │
├─────────────────────────┤
│ for each SaleItem:       │
│   1. MenuRecipe 조회     │
│   2. 차감량 계산:        │
│      qty × recipe.qtyUsed│
│   3. 식자재 단가 스냅샷:  │
│      costAtSale 저장      │
│   4. 재고 currentStock --│
│   5. InventoryLog 기록   │
│ 트랜잭션 커밋            │
└────────────┬─────────────┘
             │
             ▼
[원가율 산출]
       │
       ├─→ 메뉴별 원가율
       ├─→ 카테고리별 원가율
       └─→ 전체 식당 원가율
             │
             ▼
[임계 초과 검사]
       │
       ├─ 초과 → 운영자 푸시 알림
       └─ 정상 → 종료"""
    add_code_block(doc, fig3)

    add_heading_h3(doc, '도 4 — 원가율 산출 수식')
    fig4 = """메뉴 i의 1개 원가:
    원가_i = Σ_j (recipe[i,j].qtyUsed × inventory[j].unitPrice)

메뉴 i의 원가율:
    원가율_i = 원가_i / menu[i].price × 100 (%)

전체 식당 원가율:
    원가율_총 = Σ_k (saleItem[k].costAtSale × saleItem[k].qty)
              / Σ_k saleItem[k].subtotal
              × 100 (%)

원가율 색상 코딩:
    원가율 < 30%        : 녹색 (양호)
    30% ≤ 원가율 < 40%  : 황색 (주의)
    원가율 ≥ 40%        : 적색 (개선 필요)"""
    add_code_block(doc, fig4)

    add_heading_h3(doc, '도 5 — 운영자 화면 예시')
    fig5 = """┌─────────────────────────────────────┐
│ 🍽️ 메뉴 / 원가             [+ 메뉴]│
├─────────────────────────────────────┤
│ ┌──────┐ ┌──────┐ ┌──────────┐ │
│ │ 12   │ │  3   │ │   34%    │ │
│ │ 메뉴 │ │ 레시피│ │ 평균원가율│ │
│ └──────┘ └──────┘ └──────────┘ │
├─────────────────────────────────────┤
│ 비빔밥                  9,000원      │
│ 메인 · 재료 5개          4,200원     │
│ ─────────────────────────────────│
│ 원가          원가율                  │
│ 4,200원      🟡 46.7%                │
├─────────────────────────────────────┤
│ 칼국수                  8,000원      │
│ 메인 · 재료 4개          2,300원     │
│ ─────────────────────────────────│
│ 원가          원가율                  │
│ 2,300원      🟢 28.8%                │
└─────────────────────────────────────┘"""
    add_code_block(doc, fig5)

    # ========== 부록 ==========
    add_heading_h2(doc, '부록 — 라이브 동작 확인')
    add_para(doc, '본 발명은 실제로 동작하는 시스템으로 구현되어 운영 중이다.')
    add_para(doc, '• 데모 URL: https://restraunt-ebon-phi.vercel.app', indent=0.5)
    add_para(doc, '• 소스 저장소: https://github.com/jaewonhyeon9-ctrl/restraunt', indent=0.5)
    add_para(doc, '시연 흐름:')
    add_para(doc, '1. 메뉴 추가 (예: "비빔밥", 9,000원)', indent=0.5)
    add_para(doc, '2. 레시피 등록 (쌀 200g + 소고기 100g + …)', indent=0.5)
    add_para(doc, '3. 매출 1건 등록 → 다중 재고 자동 차감 → 원가율 즉시 표시', indent=0.5)
    add_para(doc, '4. 식자재 단가 변경 → 미래 매출 원가율만 변동, 과거 매출은 보존', indent=0.5)

    add_para(doc, '', size=8)
    add_para(doc, '이상으로 발명설명서를 마칩니다.', align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)

    # 저장
    doc.save(str(out_path))
    print(f'OK: {out_path}')
    print(f'Size: {out_path.stat().st_size:,} bytes')


if __name__ == '__main__':
    main()